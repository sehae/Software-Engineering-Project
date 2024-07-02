import configparser
import pandas as pd
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sqlalchemy import create_engine

# Define your userManager class here or import it as needed
from validator.user_manager import userManager

# Function to load configuration from config.ini
def load_config():
    config = configparser.ConfigParser()
    config.read('config.ini')
    return config

# Function to save configuration to config.ini
def save_config(sales_path):
    config = configparser.ConfigParser()
    config.read('config.ini')

    if sales_path:
        if 'DEFAULT' not in config:
            config['DEFAULT'] = {}
        config['DEFAULT']['sales_path'] = sales_path

    with open('config.ini', 'w') as configfile:
        config.write(configfile)

# Replace 'username', 'password', 'localhost', 'dbname' with your actual MySQL credentials and database name
engine = create_engine('mysql+pymysql://root:root@localhost/poswithinventorysystem')

# Query to join order, product, and add_on tables
query = """
SELECT 
    o.Order_ID,
    o.Date,
    o.Time,
    o.Total_Amount,
    o.Order_Type,
    CASE 
        WHEN o.Order_Type = 'Package' THEN p.Package_Name  -- Retrieve package name for Package orders
        WHEN o.Order_Type = 'Add-ons only' THEN JSON_ARRAYAGG(JSON_OBJECT('quantity', ad.quantity, 'product_id', ad.product_id))  -- Construct JSON array for Add-ons only orders
        ELSE NULL
    END AS Product_Details
FROM 
    `order` o
LEFT JOIN 
    package p ON o.Package_ID = p.Package_ID
LEFT JOIN 
    (
        SELECT 
            a.Order_ID,
            JSON_UNQUOTE(JSON_EXTRACT(ad.product_details, '$.product_id')) AS product_id,
            JSON_UNQUOTE(JSON_EXTRACT(ad.product_details, '$.quantity')) AS quantity
        FROM 
            add_on a
            CROSS JOIN JSON_TABLE(a.product_details, '$[*]' COLUMNS (
                product_details JSON PATH '$'
            )) AS ad
    ) ad ON o.Order_ID = ad.Order_ID
GROUP BY 
    o.Order_ID, o.Date, o.Time, o.Total_Amount, o.Order_Type, p.Package_Name;
"""

# Read data into a pandas DataFrame
print("Reading data from the database...")
dataframe = pd.read_sql(query, engine)
print(f"Data read from database. DataFrame shape: {dataframe.shape}")
print(dataframe.head())

# Combine Date and Time into a single DateTime column
dataframe['DateTime'] = pd.to_datetime(dataframe['Date'] + ' ' + dataframe['Time'])
print(f"Date and Time combined. DataFrame shape: {dataframe.shape}")
print(dataframe.head())

# Function to generate a daily report
def generate_daily_report():
    print("Generating daily report...")
    data = dataframe.copy()
    today = datetime.now().date()
    daily_data = data[data['DateTime'].dt.date == today]
    print(f"Daily data count: {len(daily_data)}")
    print(daily_data.head())  # Print first few rows to inspect
    return daily_data

# Function to generate a weekly report
def generate_weekly_report():
    print("Generating weekly report...")
    data = dataframe.copy()
    today = datetime.now().date()
    last_week = today - timedelta(days=7)
    weekly_data = data[(data['DateTime'].dt.date >= last_week) & (data['DateTime'].dt.date <= today)]
    print(f"Weekly data count: {len(weekly_data)}")
    print(weekly_data.head())  # Print first few rows to inspect
    return weekly_data

# Function to generate a monthly report
def generate_monthly_report():
    print("Generating monthly report...")
    data = dataframe.copy()
    today = datetime.now().date()
    last_month = today - timedelta(days=30)
    monthly_data = data[(data['DateTime'].dt.date >= last_month) & (data['DateTime'].dt.date <= today)]
    print(f"Monthly data count: {len(monthly_data)}")
    print(monthly_data.head())  # Print first few rows to inspect
    return monthly_data

# Function to calculate peak hours
def calculate_peak_hours(data):
    print("Calculating peak hours...")
    data['Hour'] = data['DateTime'].dt.hour
    peak_hours = data.groupby('Hour')['Total_Amount'].sum()
    print("Peak hours calculated.")
    print(peak_hours)
    return peak_hours

# Function to calculate peak weeks (most sales on specific days like Tuesdays)
def calculate_peak_weeks(data):
    print("Calculating peak weeks...")
    data['Weekday'] = data['DateTime'].dt.weekday_name
    peak_weeks = data.groupby('Weekday')['Total_Amount'].sum()
    print("Peak weeks calculated.")
    print(peak_weeks)
    return peak_weeks

# Function to calculate average peaks of sales within the month
def calculate_average_peaks(data):
    print("Calculating average peaks...")
    data['Day'] = data['DateTime'].dt.day
    average_peaks = data.groupby('Day')['Total_Amount'].mean()
    print("Average peaks calculated.")
    print(average_peaks)
    return average_peaks

# Function to identify mostly sold products within a specified time frame
def identify_mostly_sold_products(data, time_frame='day'):
    print(f"Identifying mostly sold products for {time_frame}...")
    if time_frame == 'day':
        mostly_sold = data.groupby('Product_Details').size().nlargest(5)  # Top 5 sold products within the day
    elif time_frame == 'week':
        mostly_sold = data.groupby('Product_Details').size().nlargest(5)  # Top 5 sold products within the week
    elif time_frame == 'month':
        mostly_sold = data.groupby('Product_Details').size().nlargest(5)  # Top 5 sold products within the month
    print("Mostly sold products identified.")
    print(mostly_sold)
    return mostly_sold

# Function to save report to Excel
def save_trend_report_to_excel(report_data, report_type, file_path):
    print(f"Saving {report_type} report to Excel...")
    filename = f'{file_path}/{report_type}_report.xlsx'
    report_data.to_excel(filename, index=False)
    print(f"{report_type.capitalize()} report has been generated and saved to '{filename}'")

# Function to save report to Word document
def save_trend_report_to_word(report_type, file_path):
    print(f"Saving {report_type} report to Word document...")
    document = Document()

    section = document.sections[0]
    header = section.header
    month = datetime.now().strftime("%B")
    content_header = [
        "Moon Hey Hotpot and Grill",
        "848A Banawe St, Quezon City, 1114 Metro Manila",
        "0917 624 9289",
        f"Sales {report_type} Report",
        f"({month})",
    ]

    for content_h in content_header:
        header_paragraph = header.add_paragraph(content_h)
        run = header_paragraph.runs[0]

        # Center align the paragraph
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Remove spacing before and after the paragraph
        header_paragraph.paragraph_format.space_before = Pt(0)
        header_paragraph.paragraph_format.space_after = Pt(0)

    # Initialize mostly_sold variable
    mostly_sold = {}

    # Add plots to the document
    if report_type == 'daily':
        data = generate_daily_report()
        peak_hours = calculate_peak_hours(data)
        mostly_sold = identify_mostly_sold_products(data, 'day')

        # Plot peak hours
        print("Generating peak hours plot...")
        plt.figure(figsize=(10, 6))
        peak_hours.plot(kind='bar', color='blue')
        plt.title('Peak Hours')
        plt.xlabel('Hour of Day')
        plt.ylabel('Total Sales Amount')
        plt.savefig(f'{file_path}/peak_hours.png')
        plt.close()

        document.add_picture(f'{file_path}/peak_hours.png', width=Inches(6))

    elif report_type == 'weekly':
        data = generate_weekly_report()
        peak_weeks = calculate_peak_weeks(data)
        mostly_sold = identify_mostly_sold_products(data, 'week')

        # Plot peak weeks
        print("Generating peak weeks plot...")
        plt.figure(figsize=(10, 6))
        peak_weeks.plot(kind='bar', color='green')
        plt.title('Peak Weeks')
        plt.xlabel('Day of Week')
        plt.ylabel('Total Sales Amount')
        plt.savefig(f'{file_path}/peak_weeks.png')
        plt.close()

        document.add_picture(f'{file_path}/peak_weeks.png', width=Inches(6))

    elif report_type == 'monthly':
        data = generate_monthly_report()
        average_peaks = calculate_average_peaks(data)
        mostly_sold = identify_mostly_sold_products(data, 'month')

        # Plot average peaks
        print("Generating average peaks plot...")
        plt.figure(figsize=(10, 6))
        average_peaks.plot(kind='line', marker='o', color='orange')
        plt.title('Average Peaks of Sales within the Month')
        plt.xlabel('Day of Month')
        plt.ylabel('Average Sales Amount')
        plt.grid(True)
        plt.savefig(f'{file_path}/average_peaks.png')
        plt.close()

        document.add_picture(f'{file_path}/average_peaks.png', width=Inches(6))

    # Add table of mostly sold products if it's populated
    if mostly_sold:
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Count'

        for product_details, count in mostly_sold.items():
            row_cells = table.add_row().cells
            row_cells[0].text = product_details
            row_cells[1].text = str(count)

    # Footer
    footer = section.footer
    date = datetime.now().strftime("%B %d, %Y")
    time = datetime.now().strftime("%I:%M %p")
    user_manager = userManager._instance
    username = user_manager.get_current_username()
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.text = f"{date} | {time}    Created by: {username}"
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document
    filename = f'{file_path}/{report_type}_trend_report.docx'
    document.save(filename)
    print(f"{report_type.capitalize()} report has been generated and saved to '{filename}'")
